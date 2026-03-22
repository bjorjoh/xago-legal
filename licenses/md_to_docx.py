"""Convert markdown files in this directory to .docx Word documents."""

import re
import glob
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()

    # Narrow margins (1.27 cm / 0.5 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Heading styles
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Skip horizontal rules
        if re.match(r"^---+\s*$", line):
            i += 1
            continue

        # Skip empty lines
        if line.strip() == "":
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,3})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_heading(level=level)
            add_rich_text(p, text)
            i += 1
            continue

        # Table detection
        if "|" in line and i + 1 < len(lines) and re.match(r"^\|[-\s|:]+\|$", lines[i + 1].strip()):
            table_lines = [line]
            i += 1
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            add_table(doc, table_lines)
            continue

        # Simple table (no separator line) - like key-value tables
        if "|" in line and line.strip().startswith("|"):
            table_lines = [line]
            i += 1
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            add_table(doc, table_lines)
            continue

        # Blockquote
        if line.startswith(">"):
            text = re.sub(r"^>\s*", "", line)
            # Remove emoji
            text = re.sub(r"⚠️\s*", "", text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            add_rich_text(p, text)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r"^(\d+)\.\s+(.*)", line)
        if ol_match:
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, ol_match.group(2))
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r"^[-*]\s+(.*)", line)
        if ul_match:
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, ul_match.group(1))
            i += 1
            continue

        # Bold-only line (like subtitle)
        bold_line_match = re.match(r"^\*\*(.*)\*\*$", line.strip())
        if bold_line_match:
            p = doc.add_paragraph()
            run = p.add_run(bold_line_match.group(1))
            run.bold = True
            i += 1
            continue

        # Italic line (like footer note)
        italic_match = re.match(r"^\*(.*)\*$", line.strip())
        if italic_match:
            p = doc.add_paragraph()
            run = p.add_run(italic_match.group(1))
            run.italic = True
            run.font.size = Pt(9)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_rich_text(p, line)
        i += 1

    doc.save(docx_path)


def add_rich_text(paragraph, text):
    """Add text with bold/italic markdown formatting to a paragraph."""
    # Process **bold** and *italic* markers
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_table(doc, table_lines):
    """Parse markdown table lines and add a Word table."""
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip separator rows
        if all(re.match(r"^[-:\s]+$", c) for c in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                add_rich_text(p, cell_text)
                # Bold first row
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True

    doc.add_paragraph()  # spacing after table


def convert_directory(src_dir):
    """Convert all .md files in src_dir to .docx in src_dir/docx/."""
    md_files = sorted(glob.glob(os.path.join(src_dir, "*.md")))
    if not md_files:
        print(f"  No .md files in {src_dir}")
        return
    out_dir = os.path.join(src_dir, "docx")
    os.makedirs(out_dir, exist_ok=True)
    for md_path in md_files:
        basename = os.path.splitext(os.path.basename(md_path))[0]
        docx_path = os.path.join(out_dir, f"{basename}.docx")
        print(f"  {os.path.basename(md_path)} -> docx/{basename}.docx")
        parse_markdown_to_docx(md_path, docx_path)


def main():
    import sys
    if len(sys.argv) > 1:
        dirs = [os.path.abspath(d) for d in sys.argv[1:]]
    else:
        dirs = [SCRIPT_DIR]

    for d in dirs:
        print(f"Processing: {d}")
        convert_directory(d)

    print("Done!")


if __name__ == "__main__":
    main()
